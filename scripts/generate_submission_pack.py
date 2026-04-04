from __future__ import annotations

import re
import textwrap
from pathlib import Path

import markdown
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = ROOT / "FINAL_SUBMISSION.md"
DOCX_PATH = ROOT / "FINAL_SUBMISSION_READY.docx"
HTML_PATH = ROOT / "FINAL_SUBMISSION_READY.html"
SCREENSHOTS_DIR = ROOT / "screenshots"


NAVY = "#081327"
CARD = "#101E36"
TEAL = "#2FD3C4"
CYAN = "#29B4FF"
GREEN = "#45D483"
TEXT = "#F2F5FA"
MUTED = "#8FA5C2"
BORDER = "#28486B"


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, box, fill, outline=BORDER, radius=22, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, start, end, color=TEAL, width=6):
    draw.line([start, end], fill=color, width=width)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    if dx == 0 and dy == 0:
        return
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux = dx / length
    uy = dy / length
    left = (end[0] - 18 * ux - 10 * uy, end[1] - 18 * uy + 10 * ux)
    right = (end[0] - 18 * ux + 10 * uy, end[1] - 18 * uy - 10 * ux)
    draw.polygon([end, left, right], fill=color)


def center_text(draw: ImageDraw.ImageDraw, box, title, subtitle=None, title_fill=TEXT):
    x1, y1, x2, y2 = box
    title_font = font(24, bold=True)
    body_font = font(15)
    bbox = draw.multiline_textbbox((0, 0), title, font=title_font, spacing=4)
    title_w = bbox[2] - bbox[0]
    title_h = bbox[3] - bbox[1]
    draw.multiline_text(
        (x1 + (x2 - x1 - title_w) / 2, y1 + 18),
        title,
        font=title_font,
        fill=title_fill,
        spacing=4,
        align="center",
    )
    if subtitle:
        wrapped = textwrap.fill(subtitle, 28)
        bbox2 = draw.multiline_textbbox((0, 0), wrapped, font=body_font, spacing=4)
        sub_w = bbox2[2] - bbox2[0]
        draw.multiline_text(
            (x1 + (x2 - x1 - sub_w) / 2, y1 + 18 + title_h + 12),
            wrapped,
            font=body_font,
            fill=MUTED,
            spacing=4,
            align="center",
        )


def wrap_text_by_width(draw: ImageDraw.ImageDraw, text: str, font_obj, max_width: int):
    words = text.split()
    if not words:
        return ""
    lines = []
    current = words[0]
    for word in words[1:]:
        test = f"{current} {word}"
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= max_width:
            current = test
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return "\n".join(lines)


def draw_box_text(
    draw: ImageDraw.ImageDraw,
    box,
    title: str,
    subtitle: str = "",
    *,
    align: str = "center",
    title_fill: str = TEXT,
    title_size: int = 24,
    body_size: int = 15,
    padding_x: int = 24,
    top_y: int = 22,
):
    x1, y1, x2, y2 = box
    title_font = font(title_size, bold=True)
    body_font = font(body_size)
    text_width = x2 - x1 - padding_x * 2
    wrapped_title = wrap_text_by_width(draw, title, title_font, text_width)
    wrapped_subtitle = wrap_text_by_width(draw, subtitle, body_font, text_width) if subtitle else ""
    title_bbox = draw.multiline_textbbox((0, 0), wrapped_title, font=title_font, spacing=4, align=align)
    title_h = title_bbox[3] - title_bbox[1]
    if align == "center":
        title_x = x1 + (x2 - x1 - (title_bbox[2] - title_bbox[0])) / 2
    else:
        title_x = x1 + padding_x
    draw.multiline_text(
        (title_x, y1 + top_y),
        wrapped_title,
        font=title_font,
        fill=title_fill,
        spacing=4,
        align=align,
    )
    if wrapped_subtitle:
        body_bbox = draw.multiline_textbbox((0, 0), wrapped_subtitle, font=body_font, spacing=6, align=align)
        if align == "center":
            body_x = x1 + (x2 - x1 - (body_bbox[2] - body_bbox[0])) / 2
        else:
            body_x = x1 + padding_x
        draw.multiline_text(
            (body_x, y1 + top_y + title_h + 14),
            wrapped_subtitle,
            font=body_font,
            fill=MUTED,
            spacing=6,
            align=align,
        )


def box_center(box):
    x1, y1, x2, y2 = box
    return ((x1 + x2) / 2, (y1 + y2) / 2)


def route_arrow(draw: ImageDraw.ImageDraw, points, color=CYAN, width=6):
    for start, end in zip(points, points[1:]):
        draw.line([start, end], fill=color, width=width)
    start = points[-2]
    end = points[-1]
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux = dx / length
    uy = dy / length
    left = (end[0] - 18 * ux - 10 * uy, end[1] - 18 * uy + 10 * ux)
    right = (end[0] - 18 * ux + 10 * uy, end[1] - 18 * uy - 10 * ux)
    draw.polygon([end, left, right], fill=color)


def pill(draw: ImageDraw.ImageDraw, box, text: str, fill: str, outline: str, text_fill: str = TEXT):
    rounded_box(draw, box, fill=fill, outline=outline, radius=30, width=2)
    f = font(15, bold=True)
    wrapped = wrap_text_by_width(draw, text, f, box[2] - box[0] - 20)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=f, spacing=2)
    draw.multiline_text(
        (box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2, box[1] + (box[3] - box[1] - (bbox[3] - bbox[1])) / 2),
        wrapped,
        font=f,
        fill=text_fill,
        spacing=2,
        align="center",
    )


def add_canvas_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str):
    draw.text((70, 48), title, font=font(30, bold=True), fill=TEXT)
    draw.text((70, 88), subtitle, font=font(16), fill=MUTED)


def generate_architecture_diagram(output_path: Path):
    image = Image.new("RGB", (1600, 900), NAVY)
    draw = ImageDraw.Draw(image)
    add_canvas_title(draw, "NorthStar Retail Architecture", "Microservices + agentic AI control plane")

    boxes = {
        "ui": (80, 250, 320, 420),
        "gateway": (390, 250, 680, 420),
        "auth": (790, 120, 1090, 260),
        "inventory": (790, 300, 1090, 440),
        "sales": (790, 480, 1090, 620),
        "ai": (1170, 230, 1510, 450),
        "tools": (1205, 490, 1475, 575),
        "db": (430, 690, 1170, 835),
    }

    rounded_box(draw, boxes["ui"], CARD)
    rounded_box(draw, boxes["gateway"], CARD)
    rounded_box(draw, boxes["auth"], CARD)
    rounded_box(draw, boxes["inventory"], CARD)
    rounded_box(draw, boxes["sales"], CARD)
    rounded_box(draw, boxes["ai"], "#0D2134", outline="#1E7486")
    rounded_box(draw, boxes["tools"], "#0D1D31", outline="#246F8C")
    rounded_box(draw, boxes["db"], "#0B1830", outline="#20546B")

    draw_box_text(draw, boxes["ui"], "Manager UI", "React dashboard\nMobile POS", align="center", title_size=24, body_size=16)
    draw_box_text(draw, boxes["gateway"], "API Gateway", "JWT verification\nService routing", align="center", title_size=25, body_size=16)
    draw_box_text(draw, boxes["auth"], "Auth Service", "Login\nRefresh tokens\nRBAC", align="center", title_size=23, body_size=15)
    draw_box_text(draw, boxes["inventory"], "Inventory Service", "Stock\nAdjustments\nAlerts", align="center", title_size=22, body_size=15)
    draw_box_text(draw, boxes["sales"], "Sales Service", "Orders\nTax\nReporting", align="center", title_size=22, body_size=15)
    draw_box_text(draw, boxes["ai"], "AI Service", "LangGraph agent\nSentinels\nTool trace + audit", align="center", title_fill="#DDFEFF", title_size=24, body_size=16)
    draw_box_text(draw, boxes["tools"], "Retail Tools", "Inventory checks\nSales summaries\nTransfer advisor", align="center", title_size=21, body_size=14)
    draw_box_text(draw, boxes["db"], "PostgreSQL", "Shared database with auth, inventory, sales, and AI audit schemas", align="center", title_size=24, body_size=16)

    pill(draw, (1250, 170, 1440, 215), "Agentic control plane", "#0C3040", "#247B8D")
    pill(draw, (515, 640, 700, 675), "Transactional core", "#13243D", "#2D5079")

    route_arrow(draw, [(320, 335), (390, 335)], CYAN)
    route_arrow(draw, [(680, 300), (730, 300), (730, 190), (790, 190)], CYAN)
    route_arrow(draw, [(680, 335), (730, 335), (730, 370), (790, 370)], CYAN)
    route_arrow(draw, [(680, 370), (730, 370), (730, 550), (790, 550)], CYAN)

    route_arrow(draw, [(1090, 370), (1135, 370), (1135, 340), (1170, 340)], TEAL)
    route_arrow(draw, [(1340, 450), (1340, 490)], TEAL)
    route_arrow(draw, [(1205, 532), (1150, 532), (1150, 550), (1090, 550)], TEAL)
    route_arrow(draw, [(1205, 532), (1150, 532), (1150, 370), (1090, 370)], TEAL)

    route_arrow(draw, [(940, 260), (940, 640), (610, 640), (610, 690)], GREEN)
    route_arrow(draw, [(940, 440), (940, 690)], GREEN)
    route_arrow(draw, [(940, 620), (940, 640), (1070, 640), (1070, 690)], GREEN)
    route_arrow(draw, [(1340, 575), (1340, 640), (1160, 640), (1160, 690)], TEAL)

    draw.text((1165, 595), "AI reads live store data\nand writes audit events", font=font(16), fill=MUTED, spacing=4)
    image.save(output_path)


def generate_query_flow_diagram(output_path: Path):
    image = Image.new("RGB", (1600, 900), NAVY)
    draw = ImageDraw.Draw(image)
    add_canvas_title(draw, "User-Initiated Agent Flow", "How a manager question becomes an audited AI decision")

    steps = [
        ("1. Store Manager", "Asks a natural-language question in the dashboard"),
        ("2. Gateway", "Passes verified user + store context to the AI service"),
        ("3. LangGraph Agent", "Chooses the right retail tools in a ReAct loop"),
        ("4. Retail Tools", "Fetch live inventory, sales, anomaly, or transfer data"),
        ("5. Audit Trail", "Stores summary, severity, model, and tools used"),
        ("6. UI Response", "Returns answer and shows trace in Agent Activity"),
    ]

    top = 200
    left = 55
    width = 230
    gap = 18
    for idx, (title, body) in enumerate(steps):
        x1 = left + idx * (width + gap)
        x2 = x1 + width
        y1, y2 = top, top + 320
        rounded_box(draw, (x1, y1, x2, y2), CARD if idx != 2 else "#10253A", outline="#256E89")
        pill(draw, (x1 + 18, y1 + 18, x1 + 68, y1 + 60), str(idx + 1), "#0E2941", "#2A8DB7")
        draw_box_text(
            draw,
            (x1 + 10, y1 + 48, x2 - 10, y2 - 10),
            title.split(". ", 1)[1],
            body,
            align="left",
            title_size=20,
            body_size=15,
            padding_x=18,
            top_y=20,
        )
        if idx < len(steps) - 1:
            route_arrow(draw, [(x2, y1 + 160), (x2 + gap, y1 + 160)], CYAN)

    note_box = (100, 610, 1500, 740)
    rounded_box(draw, note_box, "#0D1D31", outline="#274B6C")
    draw_box_text(
        draw,
        note_box,
        "Why this matters",
        "The agent is not answering from memory. It decides which retail tools to call, uses live store data, and logs the final decision so a manager can verify what happened.",
        align="left",
        title_size=21,
        body_size=17,
        padding_x=26,
        top_y=22,
    )
    image.save(output_path)


def generate_autonomous_diagram(output_path: Path):
    image = Image.new("RGB", (1600, 900), NAVY)
    draw = ImageDraw.Draw(image)
    add_canvas_title(draw, "Autonomous Sentinel Loops", "Scheduled background scans that create proactive alerts")

    scheduler = (585, 120, 1015, 245)
    rounded_box(draw, scheduler, CARD)
    draw_box_text(draw, scheduler, "FastAPI Scheduler", "Starts background loops when the AI service boots", title_size=24, body_size=16)

    restock = (120, 330, 660, 575)
    anomaly = (940, 330, 1480, 575)
    event_log = (420, 675, 1180, 820)

    rounded_box(draw, restock, "#0E2130", outline="#2D7E92")
    rounded_box(draw, anomaly, "#0E2130", outline="#6C70C8")
    rounded_box(draw, event_log, "#0B1830", outline="#216A66")

    draw_box_text(draw, restock, "Restock Sentinel", "Runs every 10 minutes.\nScans each monitored store.\nAsks the agent about stockout risk, velocity, and transfer options.", title_size=24, body_size=16)
    draw_box_text(draw, anomaly, "Guardian Sentinel", "Runs every 5 minutes.\nScans each monitored store.\nAsks the agent to detect suspicious transactions.", title_size=24, body_size=16)
    draw_box_text(draw, event_log, "ai_agent_events", "Every result becomes a visible alert in Agent Activity with severity, time, and tools used.", title_size=23, body_size=17)

    pill(draw, (245, 595, 535, 645), "Background LLM checks", "#0D3040", "#247B8D")
    pill(draw, (745, 595, 1085, 645), "Auditable outputs", "#12263C", "#2A5668")

    route_arrow(draw, [(800, 245), (800, 285), (390, 285), (390, 330)], TEAL)
    route_arrow(draw, [(800, 245), (800, 285), (1210, 285), (1210, 330)], CYAN)
    route_arrow(draw, [(390, 575), (390, 625), (700, 625), (700, 675)], GREEN)
    route_arrow(draw, [(1210, 575), (1210, 625), (960, 625), (960, 675)], GREEN)
    image.save(output_path)


def ensure_diagrams():
    SCREENSHOTS_DIR.mkdir(exist_ok=True)
    diagrams = [
        SCREENSHOTS_DIR / "submission_architecture.png",
        SCREENSHOTS_DIR / "submission_query_flow.png",
        SCREENSHOTS_DIR / "submission_autonomous_loops.png",
    ]
    generators = [
        generate_architecture_diagram,
        generate_query_flow_diagram,
        generate_autonomous_diagram,
    ]
    for path, generator in zip(diagrams, generators):
        generator(path)
    return diagrams


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def apply_run_style(run, kind: str = "normal"):
    font_obj = run.font
    if kind == "code":
        font_obj.name = "Consolas"
        font_obj.size = Pt(10.5)
    else:
        font_obj.name = "Segoe UI"
        font_obj.size = Pt(11)


def add_inline_markdown(paragraph, text: str):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            apply_run_style(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            apply_run_style(run)
        else:
            run = paragraph.add_run(token[1:-1])
            apply_run_style(run, "code")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        apply_run_style(run)


def add_code_block(document: Document, code: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    set_cell = False
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF0F7")
    p_pr.append(shd)
    for index, line in enumerate(code.rstrip().splitlines()):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        apply_run_style(run, "code")


def add_table(document: Document, rows: list[list[str]]):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            if row_index == 0:
                set_cell_shading(cell, "DCEAFB")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(value.strip())
                run.bold = True
                apply_run_style(run)
            else:
                add_inline_markdown(paragraph, value.strip())
    document.add_paragraph()


def parse_table(lines: list[str], start_index: int):
    table_lines = []
    index = start_index
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    rows = []
    for line in table_lines:
        if set(line.replace("|", "").strip()) <= {"-"}:
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
    return rows, index


def insert_image(document: Document, image_path: Path, width: float = 6.6, caption: str | None = None):
    if not image_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    if caption:
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.name = "Segoe UI"
        run.font.size = Pt(10.5)


def build_docx():
    diagrams = ensure_diagrams()
    lines = MARKDOWN_PATH.read_text(encoding="utf-8").splitlines()

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    style = document.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NorthStar Outfitters")
    run.bold = True
    run.font.name = "Segoe UI"
    run.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Agentic Retail Operations Platform")
    run.font.name = "Segoe UI"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(33, 108, 140)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Submission Pack")
    meta_run.italic = True
    meta_run.font.name = "Segoe UI"
    meta_run.font.size = Pt(11)

    insert_image(document, SCREENSHOTS_DIR / "01_dashboard.png", width=6.7, caption="Live dashboard with AI command and audit trail")
    document.add_page_break()

    mermaid_images = iter(diagrams)
    index = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            document.add_paragraph()
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip().lower()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if language == "mermaid":
                image_path = next(mermaid_images, None)
                if image_path:
                    insert_image(document, image_path, width=6.7)
            else:
                add_code_block(document, "\n".join(code_lines))
            index += 1
            continue
        if stripped.startswith("!["):
            match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if match:
                image_rel = match.group(2).replace("./", "")
                insert_image(document, ROOT / image_rel, width=6.4, caption=match.group(1) or None)
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            if rows:
                add_table(document, rows)
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading = stripped[level:].strip()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(heading)
                run.bold = True
                run.font.name = "Segoe UI"
                run.font.size = Pt(20)
            elif level == 2:
                run = paragraph.add_run(heading)
                run.bold = True
                run.font.name = "Segoe UI"
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(17, 76, 113)
            else:
                run = paragraph.add_run(heading)
                run.bold = True
                run.font.name = "Segoe UI"
                run.font.size = Pt(12.5)
            index += 1
            continue
        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(6)
            p_pr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF5FB")
            p_pr.append(shd)
            add_inline_markdown(paragraph, stripped.lstrip(">").strip())
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, numbered.group(2))
            index += 1
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:].strip())
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, stripped)
        index += 1

    document.save(DOCX_PATH)
    return DOCX_PATH


def build_html():
    diagrams = ensure_diagrams()
    text = MARKDOWN_PATH.read_text(encoding="utf-8")

    for image_path in diagrams:
        replacement = (
            f"\n\n![{image_path.stem.replace('_', ' ').title()}](./screenshots/{image_path.name})\n\n"
        )
        text = re.sub(r"```mermaid[\s\S]*?```", replacement, text, count=1)

    html_body = markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "sane_lists"],
        output_format="html5",
    )

    css = """
    :root {
      --bg: #071324;
      --card: #0d1a2d;
      --card-2: #0f2037;
      --text: #eef4fb;
      --muted: #8297b7;
      --line: #1b3556;
      --accent: #2bb8ff;
      --teal: #29d1c2;
      --green: #4ad58b;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      font-family: "Segoe UI", Arial, sans-serif;
      color: #11243b;
      background: linear-gradient(180deg, #eef5fb 0%, #e7f0fa 100%);
    }
    .page {
      max-width: 980px;
      margin: 0 auto;
      padding: 42px 44px 64px;
    }
    .hero {
      background: radial-gradient(circle at top left, #15314d 0%, var(--bg) 58%);
      color: var(--text);
      padding: 36px 38px;
      border-radius: 24px;
      box-shadow: 0 20px 50px rgba(7, 19, 36, 0.2);
      margin-bottom: 26px;
    }
    .hero h1 {
      font-size: 38px;
      margin: 0 0 6px;
      line-height: 1.06;
    }
    .hero h2 {
      margin: 0;
      color: #b8d8ff;
      font-size: 19px;
      font-weight: 600;
    }
    .hero p {
      margin: 12px 0 0;
      color: #c6d7ef;
      line-height: 1.6;
    }
    .content {
      background: rgba(255, 255, 255, 0.76);
      border: 1px solid rgba(27, 53, 86, 0.12);
      border-radius: 24px;
      padding: 34px 36px 40px;
      backdrop-filter: blur(10px);
    }
    h1, h2, h3, h4 {
      color: #102947;
      line-height: 1.18;
      page-break-after: avoid;
    }
    h1 { font-size: 30px; margin-top: 0; }
    h2 {
      font-size: 22px;
      margin: 28px 0 10px;
      border-top: 1px solid rgba(27, 53, 86, 0.12);
      padding-top: 18px;
    }
    h3 { font-size: 17px; margin: 18px 0 8px; color: #1d4b78; }
    p, li, blockquote {
      font-size: 14.5px;
      line-height: 1.7;
      color: #1f3247;
    }
    ul, ol { padding-left: 24px; }
    strong { color: #102947; }
    code {
      font-family: Consolas, monospace;
      font-size: 0.92em;
      background: #e8f0fb;
      padding: 2px 6px;
      border-radius: 6px;
    }
    pre {
      background: #0d1a2d;
      color: #e8f1fb;
      padding: 16px;
      border-radius: 14px;
      overflow: auto;
      white-space: pre-wrap;
    }
    pre code {
      background: transparent;
      color: inherit;
      padding: 0;
    }
    blockquote {
      margin: 16px 0;
      padding: 14px 18px;
      border-left: 4px solid var(--accent);
      background: #edf6ff;
      border-radius: 10px;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      margin: 18px 0 22px;
      overflow: hidden;
      border-radius: 14px;
    }
    th {
      background: #dcecff;
      color: #102947;
      font-size: 13px;
      text-align: left;
    }
    td {
      background: rgba(255, 255, 255, 0.7);
    }
    th, td {
      border: 1px solid #d6e2f0;
      padding: 10px 12px;
      font-size: 13px;
      vertical-align: top;
      line-height: 1.5;
    }
    hr { border: none; border-top: 1px solid rgba(27, 53, 86, 0.12); margin: 24px 0; }
    img {
      width: 100%;
      border-radius: 18px;
      border: 1px solid rgba(17, 41, 71, 0.12);
      box-shadow: 0 14px 28px rgba(11, 26, 45, 0.12);
      background: white;
    }
    p img {
      margin: 18px 0 6px;
    }
    a { color: #0f6fb2; text-decoration: none; }
    @media print {
      body { background: white; }
      .page { max-width: none; padding: 16px; }
      .content { border: none; box-shadow: none; padding: 0; background: white; }
      .hero { box-shadow: none; }
      h2, h3, table, img, pre, blockquote { page-break-inside: avoid; }
    }
    """

    html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>NorthStar Submission Pack</title>
  <style>{css}</style>
</head>
<body>
  <main class="page">
    <section class="hero">
      <h1>NorthStar Outfitters</h1>
      <h2>Agentic Retail Operations Platform</h2>
      <p>Submission-ready case study pack with architecture, service flow, agentic AI explanation, screenshots, and evaluation-criteria mapping.</p>
    </section>
    <section class="content">
      {html_body}
    </section>
  </main>
</body>
</html>
"""
    HTML_PATH.write_text(html, encoding="utf-8")
    return HTML_PATH


if __name__ == "__main__":
    docx_output = build_docx()
    html_output = build_html()
    print(docx_output)
    print(html_output)
